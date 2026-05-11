#!/usr/bin/env python3
"""
Build report.docx from the project narrative (mirrors report.tex).

Usage (from repo root):
  pip install python-docx
  python build_report_docx.py

Output: report.docx next to this script.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _repo_root() -> Path:
    return Path(__file__).resolve().parent


def _add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = str(val)


def _add_figure(doc: Document, image_path: Path, caption: str, width_in: float = 6.3) -> None:
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.is_file():
        run = p_cap.add_run()
        run.add_picture(str(image_path), width=Inches(width_in))
    else:
        run = p_cap.add_run(f"[Image not found: {image_path}]")
        run.italic = True
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)
    doc.add_paragraph()


def main() -> None:
    root = _repo_root()
    doc = Document()

    title = doc.add_heading(
        "Comparative Image Colorization on Places365: "
        "CNN, Conditional GAN, and Edge-Conditioned ControlNet",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Course project report (Word export)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ----- Abstract -----
    doc.add_heading("Abstract", level=1)
    _add_para(
        doc,
        "We compare three approaches to automatic colorization on a subset of "
        "Places365: (1) a small CNN predicting CIE-Lab ab from luminance L; "
        "(2) a Pix2Pix-style conditional GAN with PatchGAN discriminator, "
        "including a two-phase hyperparameter study; and (3) fine-tuning a "
        "Stable Diffusion 1.5 ControlNet conditioned on HED soft-edge maps, "
        "with luminance-lock post-processing. At a modest training budget "
        "(4k–8k images, single 24 GB GPU), the ControlNet achieves sample "
        "PSNR up to ~26.8 dB and SSIM ~0.97 while preserving input structure. "
        "We document failure modes, an edge-detector ablation, and inference "
        "latency spanning four orders of magnitude.",
    )

    # ----- 1 Introduction -----
    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Problem statement", level=2)
    _add_para(
        doc,
        "Grayscale-to-color mapping is ill-posed: many RGB images share the same "
        "luminance. Practical systems must trade colour plausibility, fidelity to "
        "ground truth when recoverable, and structural preservation.",
    )
    doc.add_heading("1.2 Research question", level=2)
    _add_para(
        doc,
        "How do a vanilla Lab-space CNN, a conditional GAN, and a diffusion-based "
        "ControlNet compare under the same small-data, single-GPU budget—and which "
        "design choices most affect structural quality for the diffusion path?",
    )
    doc.add_heading("1.3 Related work (summary)", level=2)
    _add_para(
        doc,
        "Representative lines include: deep colorization (Iizuka et al.; Larsson et al.; "
        "Zhang et al. Colorful); user-guided colorization (Zhang et al.); Pix2Pix "
        "conditional GANs (Isola et al.); instance-aware colorization (Su et al.); "
        "Colorization Transformer (Kumar et al.); latent diffusion / Stable Diffusion "
        "(Rombach et al.); ControlNet (Zhang et al.); Palette diffusion image-to-image "
        "(Saharia et al.); DDPM (Ho et al.); classical edges (Canny; HED); Places365 "
        "(Zhou et al.); SSIM (Wang et al.); conditional GANs (Mirza & Osindero); "
        "RePaint (Lugmayr et al.).",
    )
    doc.add_heading("1.4 Gap and contributions", level=2)
    _add_para(
        doc,
        "Gap: few apples-to-apples comparisons at student-scale budgets; diffusion "
        "colorization often drifts structurally without explicit safeguards; edge "
        "preprocessors are rarely ablated for this task.",
        bold=False,
    )
    _add_para(
        doc,
        "Contributions: three-method benchmark on Places365; luminance lock for "
        "diffusion outputs; empirical ablation (from-scratch vs. pretrained ControlNet; "
        "grayscale vs. Canny vs. HED softedge); reproducible software pins.",
    )

    # ----- 2 Data -----
    doc.add_heading("2. Data and preprocessing", level=1)
    _add_para(
        doc,
        "Dataset: Places365 train-standard, small (256×256) via torchvision. "
        "Subsets: 4,000 images for CNN/cGAN (resize to 128×128); 8,000 for ControlNet "
        "(resize to 384×384). Train/val split 90/10, seed 42. CNN/cGAN: RGB→Lab, "
        "L∈[0,1], ab normalized to ~[-1,1]. ControlNet: target RGB in [-1,1] for VAE; "
        "conditioning is HED softedge (default) or other preprocessors in ablations.",
    )

    # ----- 3 Methodology -----
    doc.add_heading("3. Methodology", level=1)
    _add_para(
        doc,
        "CNN: encoder–decoder ColorizationCNN, L1 on ab, Adam lr=1e−3. "
        "cGAN: generator G([L;z])→ab, PatchGAN D on concat(L,ab), BCE + λ·L1, "
        "two-phase grid over lr_G, lr_D, λ_L1 then 50-epoch retrain of best config. "
        "ControlNet: frozen SD-1.5 VAE/UNet/text encoder; trainable ControlNet from "
        "lllyasviel/control_v11p_sd15_softedge; latent ε-prediction MSE; optional "
        "luminance lock replaces predicted L with input L in Lab before RGB export.",
    )
    doc.add_paragraph(
        "Process diagram: see report.tex Figure 1 (TikZ) or recreate in slides."
    )

    # ----- 4 Setup -----
    doc.add_heading("4. Experimental setup", level=1)
    _add_para(
        doc,
        "Hardware: AWS g5.2xlarge, NVIDIA A10G 24 GB, AMD EPYC 7R32 (8 vCPU), 32 GiB RAM, "
        "Ubuntu 24.04. Software: Python 3.12, torch 2.4.1+cu121, torchvision 0.19.1+cu121, "
        "diffusers 0.31.0, transformers 4.46.3, controlnet_aux, opencv-python-headless, "
        "scikit-image. NVIDIA driver 535 with precompiled linux-modules-nvidia-535-server "
        "and nvidia-headless-no-dkms-535-server; nvidia-modprobe for /dev nodes.",
    )

    # ----- 5 Results -----
    doc.add_heading("5. Results", level=1)
    doc.add_heading("5.1 cGAN phase-1 grid (val L1 on ab, lower better)", level=2)
    rows = [
        ["1e-4", "1e-4", "50", "0.1041"],
        ["1e-4", "1e-4", "100", "0.0886"],
        ["1e-4", "1e-4", "200", "0.0706"],
        ["1e-4", "2e-4", "50", "0.1043"],
        ["1e-4", "2e-4", "100", "0.0802"],
        ["1e-4", "2e-4", "200", "0.0706"],
        ["2e-4", "1e-4", "50", "0.0894"],
        ["2e-4", "1e-4", "100", "0.0922"],
        ["2e-4", "1e-4", "200", "0.0715"],
        ["2e-4", "2e-4", "50", "0.0994"],
        ["2e-4", "2e-4", "100", "0.0878"],
        ["2e-4", "2e-4", "200", "0.0822"],
        ["4e-4", "1e-4", "50", "0.1002"],
        ["4e-4", "1e-4", "100", "0.0911"],
        ["4e-4", "1e-4", "200", "0.0736"],
        ["4e-4", "2e-4", "50", "0.0936"],
        ["4e-4", "2e-4", "100", "0.0940"],
        ["4e-4", "2e-4", "200", "0.0791"],
    ]
    _add_table(doc, ["lr_G", "lr_D", "lambda_L1", "best val L1"], rows)

    doc.add_heading("5.2 cGAN phase-2 (50 epochs, best phase-1 config)", level=2)
    _add_para(doc, "Best validation L1 ≈ 0.0700; final epoch 50 val L1 ≈ 0.0826.")

    doc.add_heading("5.3 ControlNet softedge (10 epochs, 8k images)", level=2)
    cn_rows = [
        ["1", "0.1369", "0.1263", "25.24", "0.966"],
        ["2", "0.1361", "0.1361", "22.41", "0.959"],
        ["3", "0.1354", "0.1302", "25.97", "0.964"],
        ["4", "0.1341", "0.1396", "26.07", "0.970"],
        ["5", "0.1324", "0.1350", "22.49", "0.950"],
        ["6", "0.1315", "0.1398", "26.80", "0.970"],
        ["7", "0.1332", "0.1271", "26.33", "0.970"],
        ["8", "0.1344", "0.1234", "26.58", "0.963"],
        ["9", "0.1320", "0.1441", "21.93", "0.949"],
        ["10", "0.1337", "0.1290", "25.29", "0.970"],
    ]
    _add_table(doc, ["Epoch", "train MSE", "val MSE", "sample PSNR", "sample SSIM"], cn_rows)
    _add_para(doc, "Best val MSE at epoch 8 (0.1234). Best sample PSNR epoch 6; best SSIM epoch 7.")

    doc.add_heading("5.4 CNN baseline", level=2)
    _add_para(
        doc,
        "Per-epoch CSV for the 30-epoch CNN run was not retained in-repo; behaviour "
        "matches high-λ_L1 cGAN (L1-dominated): validation L1 on ab in the ~0.07–0.08 "
        "range with desaturated qualitative output.",
    )

    # ----- 6 Experiments -----
    doc.add_heading("6. Further experiments", level=1)
    doc.add_heading("6.1 Robustness / OOD (structural conditioning ablation)", level=2)
    _add_para(
        doc,
        "Compared: (i) ControlNet from_unet + grayscale cond—regenerates scenes; "
        "(ii) Canny + pretrained canny head—noisy boundaries; (iii) HED softedge + "
        "pretrained softedge head—best trade-off. See Appendix A, Figure D.",
    )
    doc.add_heading("6.2 Error analysis", level=2)
    _add_para(
        doc,
        "Common failures: wrong plausible object colours (ill-posedness); skin-tone "
        "drift on tiny faces; chroma errors on specular highlights; per-epoch PSNR "
        "jitter because the scored validation mini-batch is fixed but small. "
        "Representative grids: Appendix A.",
    )
    doc.add_heading("6.3 Inference efficiency", level=2)
    _add_table(
        doc,
        ["Method", "Resolution", "Forward passes", "Latency (order)"],
        [
            ["CNN", "128²", "1", "<2 ms"],
            ["cGAN generator", "128²", "1", "<2 ms"],
            ["ControlNet + UniPC", "384²", "20", "~1.5–2.5 s"],
        ],
    )

    # ----- 7 Conclusion -----
    doc.add_heading("7. Conclusion", level=1)
    _add_para(
        doc,
        "At small budgets, pretrained softedge ControlNet + luminance lock delivers "
        "the strongest colour and structure trade-off among the three methods, at "
        "~1000× inference cost vs. regression. Pretrained init and strong edge "
        "conditioning are both necessary; raw grayscale from_unet fails.",
    )

    # ----- Appendix -----
    doc.add_page_break()
    doc.add_heading("Appendix A — Qualitative comparison figures", level=1)
    _add_para(
        doc,
        "Each image is a save_image grid. CNN/cGAN rows: grayscale | prediction | "
        "ground truth. ControlNet rows: HED softedge | prediction (luminance-locked) | "
        "ground truth. The last figure is an ablation (from_unet + grayscale at 128²).",
    )

    imgs = [
        (
            root / "runs/cnn_baseline/preview_epoch_050.png",
            "Figure A.1 — CNN baseline (epoch 50, 128²).",
        ),
        (
            root / "runs/my_cgan_study/phase2_best/preview_epoch_050.png",
            "Figure A.2 — cGAN phase-2 best (epoch 50, 128²).",
        ),
        (
            root / "runs/controlnet_softedge/preview_epoch_010.png",
            "Figure A.3 — ControlNet softedge (epoch 10, 384²).",
        ),
        (
            root / "runs/controlnet_recolor_10ep/preview_epoch_004.png",
            "Figure A.4 — Ablation: from_unet + grayscale (epoch 4, 128²).",
        ),
    ]
    for path, cap in imgs:
        _add_figure(doc, path, cap)

    doc.add_heading("Appendix B — Reproducibility", level=1)
    _add_para(
        doc,
        "See README.md and requirements-controlnet.txt in the repository. "
        "LaTeX source: report.tex. Regenerate this Word file with: python build_report_docx.py",
    )

    out = root / "report.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
